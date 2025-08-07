import streamlit as st
import sqlite3
import openai
from io import BytesIO
from datetime import datetime

# Inicialização do estado da sessão (para navegação e login simples)
if "page" not in st.session_state:
    st.session_state.page = "home"
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "user" not in st.session_state:
    st.session_state.user = None

# Credenciais de login válidas (simples, apenas um usuário e senha)
VALID_USERNAME = "sebastiao.junior"
VALID_PASSWORD = "107242"

# Configuração do banco de dados SQLite para armazenar evoluções
conn = sqlite3.connect("smartround.db")
cursor = conn.cursor()
cursor.execute("""
CREATE TABLE IF NOT EXISTS evolutions (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    patient TEXT,
    datetime TEXT,
    ward TEXT,
    note TEXT
)
""")
conn.commit()

# Função para calcular o escore APACHE II
def calculate_apacheII(temp=None, mean_bp=None, hr=None, rr=None,
                       ventilated=False, pao2=None, aado2=None, ph=None,
                       sodium=None, potassium=None, creatinine=None,
                       hct=None, wbc=None, gcs=None, age=None, chronic_health=0):
    score = 0
    missing_fields = []
    # Temperatura (°C)
    if temp is None:
        missing_fields.append("Temperatura")
    else:
        if temp > 41:
            score += 4
        elif 39 <= temp <= 40.9:
            score += 3
        elif 38.5 <= temp <= 38.9:
            score += 2
        elif 36 <= temp <= 38.4:
            score += 0
        elif 34 <= temp <= 35.9:
            score += 1
        elif 32 <= temp <= 33.9:
            score += 2
        elif 30 <= temp <= 31.9:
            score += 3
        elif temp < 30:
            score += 4
    # Pressão Arterial Média (PAM, mmHg)
    if mean_bp is None:
        missing_fields.append("Pressão Arterial (PAM)")
    else:
        if mean_bp >= 160:
            score += 4
        elif 130 <= mean_bp <= 159:
            score += 3
        elif 110 <= mean_bp <= 129:
            score += 2
        elif 70 <= mean_bp <= 109:
            score += 0
        elif 50 <= mean_bp <= 69:
            score += 2
        elif mean_bp < 50:
            score += 4
    # Frequência Cardíaca (bpm)
    if hr is None:
        missing_fields.append("Frequência Cardíaca")
    else:
        if hr >= 180:
            score += 4
        elif 140 <= hr <= 179:
            score += 3
        elif 110 <= hr <= 139:
            score += 2
        elif 70 <= hr <= 109:
            score += 0
        elif 55 <= hr <= 69:
            score += 2
        elif 40 <= hr <= 54:
            score += 3
        elif hr < 40:
            score += 4
    # Frequência Respiratória (irpm)
    if rr is None:
        missing_fields.append("Frequência Respiratória")
    else:
        if rr >= 50:
            score += 4
        elif 35 <= rr <= 49:
            score += 3
        elif 25 <= rr <= 34:
            score += 2
        elif 12 <= rr <= 24:
            score += 0
        elif 10 <= rr <= 11:
            score += 1
        elif 6 <= rr <= 9:
            score += 2
        elif rr < 6:
            score += 4
    # Oxigenação: usa gradiente A-a se ventilado (FiO2 >=0,5), senão usa PaO2
    if ventilated:
        if aado2 is None:
            missing_fields.append("Gradiente A-a O2")
        else:
            if aado2 > 500:
                score += 4
            elif 350 <= aado2 <= 499:
                score += 3
            elif 200 <= aado2 <= 349:
                score += 2
            elif aado2 < 200:
                score += 0
    else:
        if pao2 is None:
            missing_fields.append("PaO2")
        else:
            if pao2 >= 70:
                score += 0
            elif 61 <= pao2 <= 70:
                score += 1
            elif 55 <= pao2 <= 60:
                score += 3
            elif pao2 < 55:
                score += 4
    # pH arterial
    if ph is None:
        missing_fields.append("pH")
    else:
        if ph >= 7.7:
            score += 4
        elif 7.6 <= ph < 7.7:
            score += 3
        elif 7.5 <= ph < 7.6:
            score += 2
        elif 7.33 <= ph <= 7.49:
            score += 0
        elif 7.25 <= ph < 7.33:
            score += 1
        elif 7.15 <= ph < 7.25:
            score += 3
        elif ph < 7.15:
            score += 4
    # Sódio (mmol/L)
    if sodium is None:
        missing_fields.append("Sódio")
    else:
        if sodium >= 180:
            score += 4
        elif 160 <= sodium < 180:
            score += 3
        elif 155 <= sodium < 160:
            score += 2
        elif 150 <= sodium < 155:
            score += 1
        elif 130 <= sodium < 150:
            score += 0
        elif 120 <= sodium < 130:
            score += 1
        elif 111 <= sodium < 120:
            score += 2
        elif sodium <= 110:
            score += 4
    # Potássio (mmol/L)
    if potassium is None:
        missing_fields.append("Potássio")
    else:
        if potassium >= 7.0:
            score += 4
        elif 6.0 <= potassium < 7.0:
            score += 3
        elif 5.5 <= potassium < 6.0:
            score += 2
        elif 3.5 <= potassium < 5.5:
            score += 0
        elif 3.0 <= potassium < 3.5:
            score += 1
        elif 2.5 <= potassium < 3.0:
            score += 2
        elif potassium < 2.5:
            score += 4
    # Creatinina (mg/dL)
    if creatinine is None:
        missing_fields.append("Creatinina")
    else:
        if creatinine >= 3.5:
            score += 4
        elif 2.0 <= creatinine < 3.5:
            score += 3
        elif 1.5 <= creatinine < 2.0:
            score += 2
        elif 0.6 <= creatinine < 1.5:
            score += 0
        elif creatinine < 0.6:
            score += 2
    # Hematócrito (%)
    if hct is None:
        missing_fields.append("Hematócrito")
    else:
        if hct >= 60:
            score += 4
        elif 50 <= hct < 60:
            score += 2
        elif 46 <= hct < 50:
            score += 1
        elif 30 <= hct < 46:
            score += 0
        elif 20 <= hct < 30:
            score += 2
        elif hct < 20:
            score += 4
    # Leucócitos (x10^3/µL)
    if wbc is None:
        missing_fields.append("Leucócitos")
    else:
        if wbc >= 40:
            score += 4
        elif 20 <= wbc < 40:
            score += 2
        elif 15 <= wbc < 20:
            score += 1
        elif 3 <= wbc < 15:
            score += 0
        elif 1 <= wbc < 3:
            score += 3
        elif wbc < 1:
            score += 4
    # Escala de Coma de Glasgow (GCS)
    if gcs is None:
        missing_fields.append("Glasgow")
    else:
        if gcs > 15:
            gcs = 15
        score += (15 - gcs)
    # Idade (anos)
    if age is None:
        missing_fields.append("Idade")
    else:
        if age < 45:
            score += 0
        elif age <= 54:
            score += 2
        elif age <= 64:
            score += 3
        elif age <= 74:
            score += 5
        else:
            score += 6
    # Pontos por saúde crônica grave (0, 2 ou 5 conforme seleção)
    score += chronic_health
    return score, missing_fields

# Função para gerar alertas clínicos inteligentes
def generate_alerts(temp, hr, rr, spo2, wbc, gcs):
    alerts = []
    # Suspeita de Sepse (usa critérios simples de SIRS: temperatura, FC, leucócitos)
    if temp is not None and hr is not None and wbc is not None:
        if ((temp > 38 or temp < 36) and hr > 90 and (wbc > 12 or wbc < 4)):
            alerts.append("Suspeita de SEPSE (critérios SIRS positivos).")
    # Suspeita de TEP (taquicardia + hipoxemia)
    if spo2 is not None and hr is not None:
        if spo2 < 90 and hr > 100:
            alerts.append("Hipóxia e taquicardia - considerar TEP (embolia pulmonar).")
    # Suspeita de Asma Grave ou insuficiência respiratória (FR muito alta + SpO2 baixa)
    if rr is not None:
        if rr > 30 and spo2 is not None and spo2 < 92:
            alerts.append("Respiração muito rápida e SpO₂ baixa - considerar asma grave ou insuficiência respiratória.")
    # Alerta de AVC (Glasgow diminuído)
    if gcs is not None:
        if gcs < 15:
            alerts.append("Déficit neurológico (queda no nível de consciência) - considerar AVC ou outras causas.")
    return alerts

# Função para resumir exames laboratoriais e destacar alterações
def summarize_labs(labs):
    if not labs:
        return "", ""
    normals = {
        "Hemoglobina": (12, 17),
        "Hematócrito": (36, 50),
        "Leucócitos": (4, 11),
        "Plaquetas": (150, 450),
        "Sódio": (135, 145),
        "Potássio": (3.5, 5.0),
        "Creatinina": (0.6, 1.2)
    }
    full_list = []
    abn_list = []
    for lab, value in labs.items():
        if value is None or lab == "":
            continue
        val_str = f"{value}"
        if lab in normals and isinstance(value, (int, float)):
            low, high = normals[lab]
            if value < low:
                full_list.append(f"{lab}: {val_str} (baixo)")
                abn_list.append(f"{lab}: {val_str} baixo")
            elif value > high:
                full_list.append(f"{lab}: {val_str} (alto)")
                abn_list.append(f"{lab}: {val_str} alto")
            else:
                full_list.append(f"{lab}: {val_str}")
        else:
            full_list.append(f"{lab}: {val_str}")
    full_labs_str = "; ".join(full_list)
    abn_labs_str = "; ".join(abn_list)
    return full_labs_str, abn_labs_str

# Configuração da página e estilo (tela escura, ocultar rodapé padrão)
st.set_page_config(page_title="SmartRound", layout="wide")
st.markdown("""
<style>
body {
    background-color: #1E1E1E;
    color: #FFFFFF;
}
footer {visibility: hidden;}
</style>
""", unsafe_allow_html=True)
# Rodapé personalizado
st.markdown(
    "<div style='position: fixed; bottom: 0; width: 100%; text-align: center; color: gray;'>By Sebastião Almeida</div>",
    unsafe_allow_html=True
)

# Navegação entre páginas / módulos
if st.session_state.page == "home":
    st.title("SmartRound App Unificado")
    st.markdown("Bem-vindo ao SmartRound! Escolha uma opção:")
    col1, col2, col3 = st.columns([1,1,1])
    with col2:
        if st.button("SmartCheck"):
            st.session_state.page = "smartcheck"
        if st.button("SmartRound Enfermaria"):
            st.session_state.page = "login_enf"
        if st.button("SmartRound UTI"):
            st.session_state.page = "login_uti"

# Tela de Login para Enfermaria
if st.session_state.page == "login_enf":
    st.title("Login - SmartRound Enfermaria")
    user = st.text_input("Usuário")
    pwd = st.text_input("Senha", type="password")
    if st.button("Entrar"):
        if user == VALID_USERNAME and pwd == VALID_PASSWORD:
            st.session_state.logged_in = True
            st.session_state.user = user
            st.session_state.page = "enfermaria"
        else:
            st.error("Usuário ou senha incorretos. Tente novamente.")
    if st.button("Voltar"):
        st.session_state.page = "home"

# Tela de Login para UTI
if st.session_state.page == "login_uti":
    st.title("Login - SmartRound UTI")
    user = st.text_input("Usuário")
    pwd = st.text_input("Senha", type="password")
    if st.button("Entrar"):
        if user == VALID_USERNAME and pwd == VALID_PASSWORD:
            st.session_state.logged_in = True
            st.session_state.user = user
            st.session_state.page = "uti"
        else:
            st.error("Usuário ou senha incorretos. Tente novamente.")
    if st.button("Voltar"):
        st.session_state.page = "home"

# Página SmartCheck (sem necessidade de login)
if st.session_state.page == "smartcheck":
    st.title("SmartCheck - Avaliação Rápida")
    st.markdown("Preencha os dados clínicos para uma checagem rápida:")
    patient_name = st.text_input("Nome do Paciente (opcional)")
    # Campos de Sinais Vitais
    col1, col2, col3 = st.columns(3)
    with col1:
        sbp = st.number_input("PA Sistólica (mmHg)", min_value=0, max_value=300, value=120)
    with col2:
        dbp = st.number_input("PA Diastólica (mmHg)", min_value=0, max_value=200, value=80)
    with col3:
        map_val = None
        if sbp and dbp:
            map_val = (sbp + 2*dbp)/3  # cálculo da PAM
        st.write(f"PAM calculada: {map_val:.0f} mmHg" if map_val else "PAM: N/A")
    col4, col5, col6 = st.columns(3)
    with col4:
        hr = st.number_input("Frequência Cardíaca (bpm)", min_value=0, max_value=300, value=80)
    with col5:
        rr = st.number_input("Frequência Respiratória (irpm)", min_value=0, max_value=100, value=18)
    with col6:
        temp = st.number_input("Temperatura (°C)", min_value=30.0, max_value=45.0, value=36.5)
    spo2 = st.number_input("SpO₂ (%)", min_value=0, max_value=100, value=98)
    diet = st.text_input("Dieta", value="Oral livre")
    antibiotics = st.text_input("Antibióticos em uso", value="-")
    # Expansor para inserir valores de exames laboratoriais
    lab_values = {}
    with st.expander("Exames Laboratoriais"):
        wbc = st.number_input("Leucócitos (x10³/µL)", min_value=0.0, max_value=100.0, value=7.0)
        hb = st.number_input("Hemoglobina (g/dL)", min_value=0.0, max_value=25.0, value=13.0)
        hct = st.number_input("Hematócrito (%)", min_value=0.0, max_value=70.0, value=40.0)
        platelets = st.number_input("Plaquetas (x10³/µL)", min_value=0.0, max_value=1000.0, value=250.0)
        na = st.number_input("Sódio (mmol/L)", min_value=0.0, max_value=200.0, value=140.0)
        k = st.number_input("Potássio (mmol/L)", min_value=0.0, max_value=10.0, value=4.0)
        creat = st.number_input("Creatinina (mg/dL)", min_value=0.0, max_value=20.0, value=1.0)
        ph = st.number_input("pH arterial", min_value=6.5, max_value=7.8, value=7.40)
        ventilated = st.checkbox("Ventilação Mecânica (FiO₂ ≥ 50%)", value=False)
        pao2 = None
        aado2 = None
        if ventilated:
            aado2 = st.number_input("Gradiente Alvéolo-arterial O₂ (mmHg)", min_value=0, max_value=800, value=0)
        else:
            pao2 = st.number_input("PaO₂ (mmHg)", min_value=0, max_value=500, value=90)
        gcs = st.number_input("Escala de Coma de Glasgow (GCS)", min_value=3, max_value=15, value=15)
        age = st.number_input("Idade (anos)", min_value=0, max_value=120, value=50)
        chronic = st.selectbox("Doença crônica grave / Imunossupressão prévia",
                               ["Nenhuma", "Sim - pós-operatório eletivo", "Sim - não operado ou pós-op de emergência"])
        chronic_points = 0
        if chronic == "Sim - pós-operatório eletivo":
            chronic_points = 2
        elif chronic == "Sim - não operado ou pós-op de emergência":
            chronic_points = 5
        lab_values = {
            "Leucócitos": wbc,
            "Hemoglobina": hb,
            "Hematócrito": hct,
            "Plaquetas": platelets,
            "Sódio": na,
            "Potássio": k,
            "Creatinina": creat
        }
        troponin_text = st.text_input("Resultados adicionais (ex: Troponina, Lactato etc)")
    # Opção de IA médica (GPT-4)
    ai_interpret = st.checkbox("IA médica (GPT-4) para interpretação", value=False)
    api_key = ""
    if ai_interpret:
        api_key = st.text_input("OpenAI API Key", type="password")
    # Botão para gerar a evolução / resultados
    if st.button("Gerar Evolução / Resultado"):
        # Cálculo do APACHE II (não obrigatório, mas será calculado se dados presentes)
        apache_score, missing_apache = calculate_apacheII(temp=temp, mean_bp=map_val, hr=hr, rr=rr,
                                                          ventilated=ventilated, pao2=pao2, aado2=aado2, ph=ph,
                                                          sodium=na, potassium=k, creatinine=creat,
                                                          hct=hct, wbc=wbc, gcs=gcs, age=age, chronic_health=chronic_points)
        # Resumo dos exames laboratoriais
        full_labs_str, abn_labs_str = summarize_labs(lab_values)
        # Monta o texto estruturado da evolução
        note_lines = []
        note_lines.append(f"Sinais Vitais: PA {int(sbp)}/{int(dbp)} mmHg, FC {int(hr)} bpm, FR {int(rr)} irpm, Temp {temp:.1f} °C, SpO₂ {int(spo2)}%.")
        note_lines.append(f"Dieta: {diet if diet else '-'}")
        note_lines.append(f"Antibióticos: {antibiotics if antibiotics else 'Nenhum'}")
        if full_labs_str:
            note_lines.append(f"Exames: {full_labs_str}.")
        if apache_score is not None:
            if missing_apache:
                note_lines.append(f"APACHE II (parcial): {apache_score} pontos (dados faltantes para cálculo completo: {', '.join(missing_apache)}).")
            else:
                note_lines.append(f"APACHE II: {apache_score} pontos.")
        evolution_text = "\n".join(note_lines)
        # Exibe o texto da evolução gerada
        st.subheader("Evolução Gerada:")
        st.text_area("Texto da Evolução", evolution_text, height=150)
        # Se a IA estiver ativada, chama a API do OpenAI para interpretação
        if ai_interpret:
            if api_key:
                try:
                    openai.api_key = api_key
                    # use st.secrets se disponível
                    if st.secrets.get("OPENAI_API_KEY"):
                        openai.api_key = st.secrets["OPENAI_API_KEY"]
                    prompt = "Analise os seguintes dados do paciente e forneça uma interpretação clínica breve em Português:\n"
                    prompt += f"{evolution_text}\n"
                    if abn_labs_str:
                        prompt += f"Alterações laboratoriais importantes: {abn_labs_str}.\n"
                    prompt += "Quais são os possíveis problemas clínicos e sugestões de conduta?"
                    response = openai.ChatCompletion.create(
                        model="gpt-4",
                        messages=[
                            {"role": "system", "content": "Você é um médico assistente experiente fornecendo análise e recomendações."},
                            {"role": "user", "content": prompt}
                        ]
                    )
                    analysis = response.choices[0].message["content"]
                    st.subheader("Interpretação da IA:")
                    st.write(analysis)
                except Exception as e:
                    st.error("Falha na análise de IA. Verifique a chave API e tente novamente.")
            else:
                st.error("Por favor, insira a chave da API do OpenAI para usar a IA médica.")
        # Gera lista de alertas clínicos
        alerts = generate_alerts(temp=temp, hr=hr, rr=rr, spo2=spo2, wbc=wbc, gcs=gcs)
        if troponin_text:
            # Verifica menção de troponina no texto e se há número (valor)
            troponin_keywords = ["troponina", "troponin"]
            if any(word.lower() in troponin_text.lower() for word in troponin_keywords):
                if any(char.isdigit() for char in troponin_text):
                    alerts.append("Troponina relatada - avaliar possibilidade de IAM.")
        # Exibe alertas no canto superior direito (coluna pequena à direita)
        if alerts:
            colA, colB = st.columns([3,1])
            with colB:
                for alert in alerts:
                    st.error(alert)
        # Salva evolução no banco de dados (se nome do paciente fornecido)
        if patient_name:
            cursor.execute("INSERT INTO evolutions (patient, datetime, ward, note) VALUES (?, ?, ?, ?)",
                           (patient_name, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "SmartCheck", evolution_text))
            conn.commit()
            st.success("Evolução salva no banco de dados.")
        # Prepara arquivo .docx para download com o texto da evolução
        from docx import Document
        doc = Document()
        for line in note_lines:
            doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("Baixar Evolução em .docx", data=buffer, file_name="evolucao.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Página SmartRound Enfermaria (após login)
if st.session_state.page == "enfermaria" and st.session_state.logged_in:
    st.title("SmartRound - Enfermaria")
    st.markdown(f"Usuário: **{st.session_state.user}**  |  Módulo: Enfermaria")
    patient_name = st.text_input("Nome do Paciente")
    # Sinais vitais
    col1, col2, col3 = st.columns(3)
    with col1:
        sbp = st.number_input("PA Sistólica (mmHg)", min_value=0, max_value=300, value=120)
    with col2:
        dbp = st.number_input("PA Diastólica (mmHg)", min_value=0, max_value=200, value=80)
    with col3:
        map_val = None
        if sbp and dbp:
            map_val = (sbp + 2*dbp)/3
        st.write(f"PAM: {map_val:.0f} mmHg" if map_val else "PAM: N/A")
    col4, col5, col6 = st.columns(3)
    with col4:
        hr = st.number_input("Frequência Cardíaca (bpm)", min_value=0, max_value=300, value=80)
    with col5:
        rr = st.number_input("Frequência Respiratória (irpm)", min_value=0, max_value=100, value=18)
    with col6:
        temp = st.number_input("Temperatura (°C)", min_value=30.0, max_value=45.0, value=37.0)
    spo2 = st.number_input("SpO₂ (%)", min_value=0, max_value=100, value=98)
    diet = st.text_input("Dieta", value="Oral livre")
    antibiotics = st.text_input("Antibióticos", value="-")
    # Exames laboratoriais
    lab_values = {}
    with st.expander("Exames Laboratoriais"):
        wbc = st.number_input("Leucócitos (x10³/µL)", min_value=0.0, max_value=100.0, value=7.0)
        hb = st.number_input("Hemoglobina (g/dL)", min_value=0.0, max_value=25.0, value=13.0)
        hct = st.number_input("Hematócrito (%)", min_value=0.0, max_value=70.0, value=40.0)
        platelets = st.number_input("Plaquetas (x10³/µL)", min_value=0.0, max_value=1000.0, value=250.0)
        na = st.number_input("Sódio (mmol/L)", min_value=0.0, max_value=200.0, value=140.0)
        k = st.number_input("Potássio (mmol/L)", min_value=0.0, max_value=10.0, value=4.0)
        creat = st.number_input("Creatinina (mg/dL)", min_value=0.0, max_value=20.0, value=1.0)
        ph = st.number_input("pH arterial", min_value=6.5, max_value=7.8, value=7.40)
        ventilated = st.checkbox("Ventilação Mecânica (FiO₂ ≥ 50%)", value=False)
        pao2 = None
        aado2 = None
        if ventilated:
            aado2 = st.number_input("Gradiente Alvéolo-arterial O₂ (mmHg)", min_value=0, max_value=800, value=0)
        else:
            pao2 = st.number_input("PaO₂ (mmHg)", min_value=0, max_value=500, value=90)
        gcs = st.number_input("Glasgow (GCS)", min_value=3, max_value=15, value=15)
        age = st.number_input("Idade (anos)", min_value=0, max_value=120, value=50)
        chronic = st.selectbox("Doença crônica grave / Imunossupressão prévia",
                               ["Nenhuma", "Sim - pós-operatório eletivo", "Sim - não operado ou pós-op emergência"])
        chronic_points = 0
        if chronic == "Sim - pós-operatório eletivo":
            chronic_points = 2
        elif chronic == "Sim - não operado ou pós-op emergência":
            chronic_points = 5
        lab_values = {
            "Leucócitos": wbc,
            "Hemoglobina": hb,
            "Hematócrito": hct,
            "Plaquetas": platelets,
            "Sódio": na,
            "Potássio": k,
            "Creatinina": creat
        }
        troponin_text = st.text_input("Outros resultados (opcional)")
    ai_interpret = st.checkbox("IA médica (GPT-4)", value=False)
    api_key = ""
    if ai_interpret:
        api_key = st.text_input("OpenAI API Key", type="password")
    if st.button("Gerar Evolução"):
        apache_score, missing_apache = calculate_apacheII(temp=temp, mean_bp=map_val, hr=hr, rr=rr,
                                                          ventilated=ventilated, pao2=pao2, aado2=aado2, ph=ph,
                                                          sodium=na, potassium=k, creatinine=creat,
                                                          hct=hct, wbc=wbc, gcs=gcs, age=age, chronic_health=chronic_points)
        full_labs_str, abn_labs_str = summarize_labs(lab_values)
        note_lines = []
        note_lines.append(f"Sinais Vitais: PA {int(sbp)}/{int(dbp)} mmHg, FC {int(hr)} bpm, FR {int(rr)} irpm, Temp {temp:.1f} °C, SpO₂ {int(spo2)}%.")
        note_lines.append(f"Dieta: {diet if diet else '-'}")
        note_lines.append(f"Antibióticos: {antibiotics if antibiotics else 'Nenhum'}")
        if full_labs_str:
            note_lines.append(f"Exames: {full_labs_str}.")
        if apache_score is not None:
            if missing_apache:
                note_lines.append(f"APACHE II (parcial): {apache_score} pontos (faltando: {', '.join(missing_apache)}).")
            else:
                note_lines.append(f"APACHE II: {apache_score} pontos.")
        evolution_text = "\n".join(note_lines)
        st.subheader("Evolução:")
        st.text_area("Texto da Evolução", evolution_text, height=150)
        if ai_interpret:
            if api_key:
                try:
                    openai.api_key = api_key
                    if st.secrets.get("OPENAI_API_KEY"):
                        openai.api_key = st.secrets["OPENAI_API_KEY"]
                    prompt = "Interprete os dados clínicos a seguir:\n"
                    prompt += f"{evolution_text}\n"
                    if abn_labs_str:
                        prompt += f"Dados laboratoriais alterados: {abn_labs_str}.\n"
                    prompt += "Dê uma análise breve e possíveis recomendações."
                    response = openai.ChatCompletion.create(
                        model="gpt-4",
                        messages=[
                            {"role": "system", "content": "Você é um médico experiente fornecendo análise."},
                            {"role": "user", "content": prompt}
                        ]
                    )
                    analysis = response.choices[0].message["content"]
                    st.subheader("Interpretação da IA:")
                    st.write(analysis)
                except Exception as e:
                    st.error("Não foi possível obter a interpretação da IA.")
            else:
                st.error("Chave API do OpenAI não fornecida.")
        alerts = generate_alerts(temp=temp, hr=hr, rr=rr, spo2=spo2, wbc=wbc, gcs=gcs)
        if troponin_text:
            if "troponin" in troponin_text.lower() or "troponina" in troponin_text.lower():
                alerts.append("Possível elevação de troponina - investigar IAM.")
        if alerts:
            colA, colB = st.columns([3,1])
            with colB:
                for alert in alerts:
                    st.error(alert)
        if patient_name:
            cursor.execute("INSERT INTO evolutions (patient, datetime, ward, note) VALUES (?, ?, ?, ?)",
                           (patient_name, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "Enfermaria", evolution_text))
            conn.commit()
            st.success("Evolução salva no banco de dados.")
        from docx import Document
        doc = Document()
        for line in note_lines:
            doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("Baixar .docx", data=buffer, file_name="evolucao.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Página SmartRound UTI (após login)
if st.session_state.page == "uti" and st.session_state.logged_in:
    st.title("SmartRound - UTI")
    st.markdown(f"Usuário: **{st.session_state.user}**  |  Módulo: UTI")
    patient_name = st.text_input("Nome do Paciente")
    col1, col2, col3 = st.columns(3)
    with col1:
        sbp = st.number_input("PA Sistólica (mmHg)", min_value=0, max_value=300, value=120)
    with col2:
        dbp = st.number_input("PA Diastólica (mmHg)", min_value=0, max_value=200, value=80)
    with col3:
        map_val = None
        if sbp and dbp:
            map_val = (sbp + 2*dbp)/3
        st.write(f"PAM: {map_val:.0f} mmHg" if map_val else "PAM: N/A")
    col4, col5, col6 = st.columns(3)
    with col4:
        hr = st.number_input("FC (bpm)", min_value=0, max_value=300, value=85)
    with col5:
        rr = st.number_input("FR (irpm)", min_value=0, max_value=100, value=18)
    with col6:
        temp = st.number_input("Temperatura (°C)", min_value=30.0, max_value=45.0, value=37.0)
    spo2 = st.number_input("SpO₂ (%)", min_value=0, max_value=100, value=95)
    diet = st.text_input("Dieta", value="SVL")
    antibiotics = st.text_input("Antibióticos", value="-")
    lab_values = {}
    with st.expander("Exames Laboratoriais"):
        wbc = st.number_input("Leucócitos (x10³/µL)", min_value=0.0, max_value=100.0, value=12.0)
        hb = st.number_input("Hemoglobina (g/dL)", min_value=0.0, max_value=25.0, value=11.0)
        hct = st.number_input("Hematócrito (%)", min_value=0.0, max_value=70.0, value=33.0)
        platelets = st.number_input("Plaquetas (x10³/µL)", min_value=0.0, max_value=1000.0, value=200.0)
        na = st.number_input("Sódio (mmol/L)", min_value=0.0, max_value=200.0, value=135.0)
        k = st.number_input("Potássio (mmol/L)", min_value=0.0, max_value=10.0, value=4.5)
        creat = st.number_input("Creatinina (mg/dL)", min_value=0.0, max_value=20.0, value=1.2)
        ph = st.number_input("pH arterial", min_value=6.5, max_value=7.8, value=7.30)
        ventilated = st.checkbox("Ventilação Mecânica (FiO₂ ≥ 50%)", value=True)
        pao2 = None
        aado2 = None
        if ventilated:
            aado2 = st.number_input("Gradiente A-a O₂ (mmHg)", min_value=0, max_value=800, value=250)
        else:
            pao2 = st.number_input("PaO₂ (mmHg)", min_value=0, max_value=500, value=60)
        gcs = st.number_input("Glasgow (GCS)", min_value=3, max_value=15, value=10)
        age = st.number_input("Idade (anos)", min_value=0, max_value=120, value=65)
        chronic = st.selectbox("Doença crônica grave / Imunossupressão",
                               ["Nenhuma", "Sim - pós-operatório eletivo", "Sim - não operado/emergência"])
        chronic_points = 0
        if chronic == "Sim - pós-operatório eletivo":
            chronic_points = 2
        elif chronic == "Sim - não operado/emergência":
            chronic_points = 5
        lab_values = {
            "Leucócitos": wbc,
            "Hemoglobina": hb,
            "Hematócrito": hct,
            "Plaquetas": platelets,
            "Sódio": na,
            "Potássio": k,
            "Creatinina": creat
        }
        troponin_text = st.text_input("Outros exames", value="")
    ai_interpret = st.checkbox("IA médica (GPT-4)", value=False)
    api_key = ""
    if ai_interpret:
        api_key = st.text_input("OpenAI API Key", type="password")
    if st.button("Gerar Evolução"):
        apache_score, missing_apache = calculate_apacheII(temp=temp, mean_bp=map_val, hr=hr, rr=rr,
                                                          ventilated=ventilated, pao2=pao2, aado2=aado2, ph=ph,
                                                          sodium=na, potassium=k, creatinine=creat,
                                                          hct=hct, wbc=wbc, gcs=gcs, age=age, chronic_health=chronic_points)
        full_labs_str, abn_labs_str = summarize_labs(lab_values)
        note_lines = []
        note_lines.append(f"Sinais Vitais: PA {int(sbp)}/{int(dbp)} mmHg, FC {int(hr)} bpm, FR {int(rr)} irpm, Temp {temp:.1f} °C, SpO₂ {int(spo2)}%.")
        note_lines.append(f"Dieta: {diet if diet else '-'}")
        note_lines.append(f"Antibióticos: {antibiotics if antibiotics else 'Nenhum'}")
        if full_labs_str:
            note_lines.append(f"Exames: {full_labs_str}.")
        if apache_score is not None:
            if missing_apache:
                note_lines.append(f"APACHE II (parcial): {apache_score} pts (faltam: {', '.join(missing_apache)}).")
            else:
                note_lines.append(f"APACHE II: {apache_score} pts.")
        evolution_text = "\n".join(note_lines)
        st.subheader("Evolução:")
        st.text_area("Texto da Evolução", evolution_text, height=150)
        if ai_interpret:
            if api_key:
                try:
                    openai.api_key = api_key
                    if st.secrets.get("OPENAI_API_KEY"):
                        openai.api_key = st.secrets["OPENAI_API_KEY"]
                    prompt = "Dados do paciente:\n"
                    prompt += f"{evolution_text}\n"
                    if abn_labs_str:
                        prompt += f"Exames alterados: {abn_labs_str}.\n"
                    prompt += "Análise e conduta em poucas frases."
                    response = openai.ChatCompletion.create(
                        model="gpt-4",
                        messages=[
                            {"role": "system", "content": "Você é um especialista em medicina crítica fornecendo recomendações."},
                            {"role": "user", "content": prompt}
                        ]
                    )
                    analysis = response.choices[0].message["content"]
                    st.subheader("Interpretação da IA:")
                    st.write(analysis)
                except Exception as e:
                    st.error("Interpretação da IA não disponível.")
            else:
                st.error("Chave API não fornecida.")
        alerts = generate_alerts(temp=temp, hr=hr, rr=rr, spo2=spo2, wbc=wbc, gcs=gcs)
        if troponin_text:
            if "troponin" in troponin_text.lower() or "troponina" in troponin_text.lower():
                alerts.append("Troponina elevada - possível IAM.")
        if alerts:
            colA, colB = st.columns([3,1])
            with colB:
                for alert in alerts:
                    st.error(alert)
        if patient_name:
            cursor.execute("INSERT INTO evolutions (patient, datetime, ward, note) VALUES (?, ?, ?, ?)",
                           (patient_name, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "UTI", evolution_text))
            conn.commit()
            st.success("Evolução salva no banco de dados.")
        from docx import Document
        doc = Document()
        for line in note_lines:
            doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("Baixar .docx", data=buffer, file_name="evolucao.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
